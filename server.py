"""
Attune — backend
================
Holds your OpenRouter API key (never exposed to the browser), reads the
uploaded photo or text, has the AI analyze the CONCEPTS, and returns
brand-new questions verified by a second AI pass.

Run locally:
    pip install -r requirements.txt
    export OPENROUTER_API_KEY=sk-or-...        (Windows: set OPENROUTER_API_KEY=sk-or-...)
    uvicorn server:app --reload
    open http://127.0.0.1:8000
"""

import os, io, json, re, base64, random, time
from typing import List, Optional
import urllib.parse

import requests
from fastapi import FastAPI, UploadFile, File, Form, HTTPException
from fastapi.staticfiles import StaticFiles
from fastapi.responses import FileResponse
from pydantic import BaseModel

API_KEY = os.environ.get("OPENROUTER_API_KEY", "").strip()
OR_URL = "https://openrouter.ai/api/v1"

app = FastAPI(title="Attune")


# ---------------------------------------------------------------- prompts
GENERATE_PROMPT = """You are an expert teacher and assessment designer building a quiz for a student.

You will receive study material: text, an image of a page, a scanned document, or even an
existing worksheet/quiz (e.g. matching columns, fill-in-the-blanks, or multiple-choice
questions). Your job is NOT to reproduce that material — it is to teach the same ideas with
FRESH questions the student has never seen.

WORK IN THIS EXACT ORDER:

STEP 1 - READ AND UNDERSTAND:
- What subject/topic is this? Be specific (e.g. "algebraic identities: expanding (a+b)^2",
  or "geography: physical landforms and their definitions" — not just "math" or "geography").
- List every distinct CONCEPT, rule, fact, definition, formula or idea the material teaches
  or tests. If the material is ITSELF a set of questions (matching, blanks, MCQs), look
  THROUGH each item to the underlying concept it is really testing, and list that concept.

STEP 2 - DESIGN NEW QUESTIONS ON THOSE SAME CONCEPTS:
- Write at least one new question for EVERY concept you listed, then add deeper questions
  that combine or apply concepts. Aim for about {n} questions, but produce MORE if the
  material is rich — cover everything it teaches. Never pad with filler if it is thin.
- Each question must test the SAME concept as the source, but be a genuinely DIFFERENT
  question: new wording, new numbers, new examples, a new scenario or a reversed direction
  (e.g. if the source matched "harbor -> definition", ask the definition and offer terms,
  or ask which term does NOT fit).
- ABSOLUTELY FORBIDDEN: copying a sentence from the material and deleting a word
  (fill-in-the-blank of the original), or repeating one of the material's own questions with
  only tiny edits. If a student memorised the input, your questions must still make them
  think. Test UNDERSTANDING of the concept, not recall of the exact wording.
- Difficulty ladder: start easy (direct recognition/application), build to hard (apply the
  concept to an unfamiliar situation, or combine two concepts).
- For math/formula material: invent NEW problems using the same rules, and SOLVE each one
  step by step in your reasoning so the answer is guaranteed correct.
- If reading an IMAGE or SCAN: first transcribe the key facts/definitions/formulas, and
  state your reading of any ambiguous symbol, before writing questions.
- Keep each question SHORT (it appears inside a game, on one line if possible).
- Use plain text only — NO LaTeX or markdown symbols (write "->" not "\\rightarrow", "x^2"
  not "$x^2$").
- 4 options each, each option short. Wrong options must be PLAUSIBLE student errors
  (sign slips, off-by-one, common misconceptions, closely-related terms) - never silly.
- Exactly ONE correct option.

STEP 3 - OUTPUT:
After your reasoning, output the quiz inside <QUIZ> tags, exactly:
<QUIZ>
{{"topic":"specific topic label","questions":[
 {{"prompt":"...","options":["...","...","...","..."],"answer":0,
  "explain":"one short sentence teaching why the correct answer is right",
  "difficulty":"easy"}}
]}}
</QUIZ>
"answer" is the 0-3 index of the correct option. Output as many question objects as the
material supports. No markdown fences inside the tags."""

VERIFY_PROMPT = """You are a meticulous exam checker. Below is a quiz in JSON. For EACH question,
independently solve it yourself from scratch WITHOUT trusting the marked answer.

- If the marked answer is correct: keep the question unchanged.
- If the marked answer is WRONG: fix the "answer" index to the truly correct option and
  update "explain".
- If a question is ambiguous, has two defensible answers, or is broken: REWRITE it into a
  clean question on the same concept with exactly one correct answer.
- Keep the same number of questions and the same JSON shape.

Show brief checking work, then output the corrected quiz inside <QUIZ> tags in the same
JSON shape. No markdown fences inside the tags.

QUIZ TO CHECK:
"""


# ---------------------------------------------------------------- model discovery
_free_vision: List[str] = []
_free_text: List[str] = []
_picked = {"vision": None, "text": None}


def discover_models():
    """Free models on OpenRouter rotate, so ask which are live right now."""
    global _free_vision, _free_text
    try:
        r = requests.get(f"{OR_URL}/models", timeout=30)
        r.raise_for_status()
        vision, text = [], []
        for m in r.json().get("data", []):
            mid = m.get("id", "")
            if not mid.endswith(":free"):
                continue
            mods = (m.get("architecture") or {}).get("input_modalities") or []
            (vision if "image" in mods else text).append(mid)
        _free_vision, _free_text = vision, text
    except Exception as e:
        print("model discovery failed:", e)
        _free_vision, _free_text = [], []
    return _free_vision, _free_text


def _rank(m: str) -> int:
    s = 0
    if "gemma" in m: s -= 5
    if any(k in m for k in ("31b", "70b", "120b")): s -= 3
    if any(k in m for k in ("nano", "mini", "small")): s += 3
    return s


def pick_model(kind: str) -> str:
    if _picked[kind]:
        return _picked[kind]
    if not _free_vision and not _free_text:
        discover_models()
    pool = _free_vision if kind == "vision" else (_free_text + _free_vision)
    if not pool:
        raise HTTPException(503, "No free models available from OpenRouter right now.")
    for m in sorted(pool, key=_rank)[:6]:
        try:
            _chat([{"role": "user", "content": "ping"}], m, 0.0, max_tokens=5, tries=1)
            _picked[kind] = m
            print(f"using {kind} model: {m}")
            return m
        except Exception:
            continue
    raise HTTPException(503, f"No usable free {kind} model right now.")


# ---------------------------------------------------------------- AI call
def _chat(messages, model, temperature, max_tokens=3000, tries=3):
    if not API_KEY:
        raise HTTPException(500, "OPENROUTER_API_KEY is not set on the server.")
    last = None
    for attempt in range(tries):
        try:
            r = requests.post(
                f"{OR_URL}/chat/completions",
                headers={"Authorization": f"Bearer {API_KEY}",
                         "Content-Type": "application/json"},
                json={"model": model, "messages": messages,
                      "temperature": temperature, "max_tokens": max_tokens},
                timeout=120,
            )
            if r.status_code != 200:
                raise RuntimeError(f"{r.status_code}: {r.text[:200]}")
            return r.json()["choices"][0]["message"]["content"]
        except Exception as e:
            last = e
            if attempt == tries - 1:
                raise
            time.sleep(2 * (attempt + 1))
    raise RuntimeError(str(last))


# ---------------------------------------------------------------- parsing
def _extract_quiz(text: str) -> dict:
    m = re.search(r"<QUIZ>(.*?)</QUIZ>", text, re.S)
    payload = m.group(1) if m else text
    payload = payload.replace("```json", "").replace("```", "").strip()
    a, z = payload.find("{"), payload.rfind("}")
    if a != -1 and z != -1:
        payload = payload[a:z + 1]
    try:
        return json.loads(payload)
    except (json.JSONDecodeError, ValueError):
        # Model didn't return usable JSON (e.g. an unreadable photo). Fall through
        # to an empty quiz so callers raise the friendly "couldn't build" message.
        return {}


_LATEX = {r"\rightarrow": "->", r"\to": "->", r"\times": "x", r"\cdot": "*",
          r"\div": "/", r"\leq": "<=", r"\geq": ">=", r"\neq": "!=",
          r"\pm": "+/-", r"\degree": "deg", r"\%": "%"}


def _plain(s: str) -> str:
    """Strip stray LaTeX/markdown so questions read as plain text in the game."""
    s = str(s)
    for k, v in _LATEX.items():
        s = s.replace(k, v)
    s = re.sub(r"\$+", "", s)            # $...$ math delimiters
    s = re.sub(r"\\[a-zA-Z]+", "", s)    # leftover \commands
    s = s.replace("\\", "").replace("**", "").replace("`", "")
    return re.sub(r"\s+", " ", s).strip()


def _sanitize(quiz: dict) -> dict:
    """Dedupe options, guarantee exactly one valid answer, shuffle positions."""
    clean = []
    for q in quiz.get("questions", []):
        if not (q.get("prompt") and isinstance(q.get("options"), list)):
            continue
        raw = [_plain(o) for o in q["options"] if _plain(o)]
        idx = q.get("answer", 0)
        if not isinstance(idx, int) or not (0 <= idx < len(raw)):
            continue
        correct = raw[idx]
        opts = list(dict.fromkeys(raw))          # dedupe, keep order
        if len(opts) < 3 or correct not in opts:
            continue
        opts = opts[:4]
        if correct not in opts:
            opts[-1] = correct
        random.shuffle(opts)
        clean.append({
            "prompt": _plain(q["prompt"]),
            "options": opts,
            "answer": opts.index(correct),
            "explain": _plain(q.get("explain", ""))[:220],
            "difficulty": q.get("difficulty", "medium"),
        })
    return {"topic": quiz.get("topic", "Practice"), "questions": clean}


def generate(user_content, kind: str, n: int) -> dict:
    # More questions => more reasoning + output tokens. Scale the budget with n.
    budget = min(16000, 3000 + n * 350)
    model = pick_model(kind)
    raw = _chat([{"role": "user", "content": user_content}], model, 0.7, max_tokens=budget)
    draft = _sanitize(_extract_quiz(raw))
    if not draft["questions"]:
        raise HTTPException(422, "Couldn't build questions from that. Try clearer material.")

    # second pass: re-solve everything and fix wrong answer keys
    try:
        tmodel = pick_model("text")
        raw2 = _chat([{"role": "user",
                       "content": VERIFY_PROMPT + json.dumps(draft, ensure_ascii=False)}],
                     tmodel, 0.0, max_tokens=budget)
        verified = _sanitize(_extract_quiz(raw2))
        if len(verified["questions"]) >= max(3, len(draft["questions"]) // 2):
            return verified
    except Exception as e:
        print("verification skipped:", e)
    return draft


# ---------------------------------------------------------------- reading files
MAX_PAGES = 25          # cap PDF pages / total images we send to the model
N_CAP = 40              # the most questions we ever ask for


def _img_part(mime: str, data: bytes) -> dict:
    b64 = base64.b64encode(data).decode()
    return {"type": "image_url", "image_url": {"url": f"data:{mime};base64,{b64}"}}


def _read_pdf(data: bytes):
    """Return (text, [image_parts]). Text pages are read directly; scanned/image-only
    pages are rendered to PNG so the vision model can read them."""
    import fitz
    texts, images = [], []
    with fitz.open(stream=data, filetype="pdf") as doc:
        for i, page in enumerate(doc):
            if i >= MAX_PAGES:
                break
            t = page.get_text().strip()
            if len(t) >= 40:
                texts.append(t)
            else:
                pix = page.get_pixmap(dpi=140)
                images.append(_img_part("image/png", pix.tobytes("png")))
    return "\n\n".join(texts), images


def _read_docx(data: bytes) -> str:
    import docx
    d = docx.Document(io.BytesIO(data))
    lines = [p.text for p in d.paragraphs if p.text.strip()]
    for tbl in d.tables:                       # matching worksheets often live in tables
        for row in tbl.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                lines.append("  |  ".join(cells))
    return "\n".join(lines)


def _read_upload(name: str, mime: str, data: bytes):
    """Turn one uploaded file into (text, [image_parts])."""
    low = (name or "").lower()
    mime = mime or ""
    if mime.startswith("image/") or low.endswith((".png", ".jpg", ".jpeg", ".webp", ".gif")):
        return "", [_img_part(mime or "image/jpeg", data)]
    if mime == "application/pdf" or low.endswith(".pdf"):
        return _read_pdf(data)
    if low.endswith(".docx"):
        return _read_docx(data), []
    if mime.startswith("text/") or low.endswith((".txt", ".md", ".csv")):
        return data.decode("utf-8", "ignore"), []
    # last resort: try to read it as text
    try:
        return data.decode("utf-8"), []
    except Exception:
        raise HTTPException(400, f"Unsupported file type: {name or 'file'}")


# ---------------------------------------------------------------- web research
UA = {"User-Agent": "Attune/1.0 (educational quiz builder)"}


def _get_json(url, timeout=15):
    r = requests.get(url, timeout=timeout, headers=UA)
    r.raise_for_status()
    return r.json()


def _search_query(text: str) -> str:
    """Ask a fast model for a concise web-search query naming the real topic."""
    try:
        model = pick_model("text")
        q = _chat([{"role": "user", "content":
            "From the study material below, output ONLY a concise web-search query naming the main "
            "topic. If it is a book, work, person or event, give the exact title/name (add the author "
            "if it is a book). No quotes, no explanation, max 8 words.\n\nMATERIAL:\n" + text[:1200]}],
            model, 0.0, max_tokens=40)
        q = (q or "").strip().splitlines()[0].strip().strip('"').strip()
        return q[:120] or text[:80]
    except Exception as e:
        print("query distil failed:", e)
        return text[:80]


def _wikipedia(query: str, max_articles=2, chars=3500):
    out = []
    try:
        s = _get_json("https://en.wikipedia.org/w/api.php?" + urllib.parse.urlencode({
            "action": "query", "list": "search", "srsearch": query,
            "format": "json", "srlimit": max_articles}))
        titles = [h["title"] for h in s.get("query", {}).get("search", [])][:max_articles]
        for title in titles:
            e = _get_json("https://en.wikipedia.org/w/api.php?" + urllib.parse.urlencode({
                "action": "query", "prop": "extracts", "explaintext": 1, "redirects": 1,
                "titles": title, "format": "json"}))
            for p in e.get("query", {}).get("pages", {}).values():
                ex = (p.get("extract") or "").strip()
                if len(ex) > 60:
                    out.append(f"## {title}\n{ex[:chars]}")
    except Exception as e:
        print("wikipedia lookup failed:", e)
    return out


def _duckduckgo(query: str):
    try:
        d = _get_json("https://api.duckduckgo.com/?" + urllib.parse.urlencode({
            "q": query, "format": "json", "no_html": 1, "skip_disambig": 1}))
        parts = []
        if d.get("AbstractText"):
            parts.append(d["AbstractText"])
        for t in d.get("RelatedTopics", [])[:6]:
            if isinstance(t, dict) and t.get("Text"):
                parts.append(t["Text"])
        return "\n".join(parts)[:1500]
    except Exception as e:
        print("duckduckgo lookup failed:", e)
        return ""


def research(text: str):
    """Look the topic up online and return (enriched_material, topic_query, found?)."""
    query = _search_query(text)
    blocks = _wikipedia(query)
    ddg = _duckduckgo(query)
    parts = []
    if blocks:
        parts.append("\n\n".join(blocks))
    if ddg:
        parts.append("Web summary:\n" + ddg)
    if not parts:
        return text, query, False
    enriched = (f"TOPIC LOOKED UP ONLINE: {query}\n\n" + "\n\n".join(parts) +
                "\n\n---\nORIGINAL INPUT FROM THE STUDENT (may be brief):\n" + text)
    return enriched, query, True


# ---------------------------------------------------------------- API
class TextIn(BaseModel):
    text: str
    n: int = 20
    research: bool = True


@app.get("/api/health")
def health():
    v, t = (_free_vision, _free_text) if (_free_vision or _free_text) else discover_models()
    return {"key_set": bool(API_KEY), "free_vision": len(v), "free_text": len(t)}


@app.post("/api/quiz/text")
def quiz_from_text(body: TextIn):
    text = body.text.strip()
    if len(text) < 2:
        raise HTTPException(400, "Type a topic, a book name, or some text.")
    # Look the topic up online first — this lets even a book title or a one-line
    # tidbit become a real quiz. Falls back to the raw text if nothing is found.
    if body.research:
        text, _topic, _found = research(text)
    elif len(text) < 30:
        raise HTTPException(400, "Add a bit more text, or turn on 'look it up online'.")
    n = max(3, min(N_CAP, body.n))
    content = GENERATE_PROMPT.format(n=n) + "\n\nSTUDY MATERIAL:\n\n" + text
    return generate(content, "text", n)


@app.post("/api/quiz/file")
async def quiz_from_files(files: List[UploadFile] = File(...), n: int = Form(20),
                          research_flag: bool = Form(True)):
    """Photos AND documents (PDF, Word, txt). Reads/scans them all, then builds one quiz."""
    if not files:
        raise HTTPException(400, "Add at least one file.")
    texts, images, total = [], [], 0
    for f in files:
        data = await f.read()
        total += len(data)
        if total > 40 * 1024 * 1024:
            raise HTTPException(413, "Those files are too large together (max ~40 MB).")
        t, imgs = _read_upload(f.filename, f.content_type or "", data)
        if t.strip():
            texts.append(t.strip())
        images.extend(imgs)
    images = images[:MAX_PAGES]
    if not texts and not images:
        raise HTTPException(422, "Couldn't read anything from those files. Try clearer material.")

    n = max(3, min(N_CAP, n))
    text_blob = "\n\n".join(texts).strip()
    # Thin text upload (e.g. a title page or a short note) and no images -> look it up online.
    if research_flag and not images and 0 < len(text_blob) < 600:
        text_blob, _topic, _found = research(text_blob)
    if images:
        lead = "\n\nSTUDY MATERIAL is below. Read EVERYTHING — "
        lead += (f"{len(images)} page image(s)" if images else "")
        lead += (" plus this text:\n\n" + text_blob) if text_blob else " — then make the questions."
        content = [{"type": "text", "text": GENERATE_PROMPT.format(n=n) + lead}] + images
        return generate(content, "vision", n)
    content = GENERATE_PROMPT.format(n=n) + "\n\nSTUDY MATERIAL:\n\n" + text_blob
    return generate(content, "text", n)


# backwards-compatible alias
@app.post("/api/quiz/image")
async def quiz_from_image(files: List[UploadFile] = File(...), n: int = Form(20)):
    return await quiz_from_files(files, n)


# ---------------------------------------------------------------- static site
app.mount("/static", StaticFiles(directory="static"), name="static")


@app.get("/")
def index():
    # The discovery session is the front door — it profiles how the child
    # thinks before any lesson begins.
    return FileResponse("static/discovery.html")


@app.get("/game")
def game():
    # The textbook-photo runner game (build questions from a page, then play).
    return FileResponse("static/index.html")
